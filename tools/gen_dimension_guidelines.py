"""Generate the CAD dimension-guidelines Word document.

One row per generated component drawing (46 of them), stating which dimensions
the drawing uses and where each value is taken from. The content mirrors the
resolution rules in backend/main.py and backend/cad/*.py — when a rule changes
there, update it here and re-run:

    python tools/gen_dimension_guidelines.py

The .docx is written to the project root and is gitignored (it is a generated
document, not source).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "CAD Drawing Dimension Guidelines.docx"

RES_BLUE = RGBColor(0x1F, 0x4E, 0x79)
HDR_FILL = "1F4E79"
BAND_FILL = "DCE6F1"

# --------------------------------------------------------------------------- #
# Component data: (name, [guideline lines], [source lines])
# --------------------------------------------------------------------------- #
PELLET_TOL = "Dia tolerance +0.0 / -0.2, thickness tolerance ±0.05 (defaults)."

COMPONENTS: list[tuple[str, list[str], list[str]]] = [
    # ---- Assembly group ---------------------------------------------------
    ("CONTAINER", [
        "Container OD — the battery diameter.",
        "Container height.",
        "Wall thickness — Table 2, by OD band: 30–50 → 0.70; 50–100 → 1.10 deep drawn / 1.00 flanged; "
        "100–150 → 1.30. The same row gives the wall tolerance and min/max.",
        "Container ID = OD − 2 × wall (derived on every render, never frozen).",
        "Container type (deep drawn / flanged), bottom radius, flange kind / position / width — chosen by the user.",
    ], [
        "OD: manual entry → Table 1 “Diameter of Battery”.",
        "Height: manual entry → Table 1 “Height of Container”.",
        "Wall: Table 2 lookup (backend/cad/tables.py) unless overridden in CAD Revision.",
        "Refuses to draw (400) if OD or height cannot be found.",
    ]),
    ("LID BLANK", [
        "Lid blank OD = Container ID − 0.05  (Table 4a).",
        "Thickness — Table 4(b), by container OD: 30–50 → 2.5; 51–80 → 4.0; 81–110 → 5.0; 111–150 → 6.5.",
        "Groove depth & width — Table 3, by container OD: 30–45 → 0.80 / 0.70; 46–100 → 1.00 / 1.00; "
        "101–150 → 1.30 / 1.30.",
        "Number of terminal holes N → spacing θ = 360° / N. Defaults to 4 with a warning.",
        "PCD of the holes = Pin PCD; if absent, cathode dia − clearance (clearance 1.0 default).",
        "Hole dia = pin dia × 2.5.",
        "Groove circle dia = cathode (stack) dia. Weld space = container wall thickness.",
        "Hole start angle 90°; edge angle, letter size and back groove are standard values, all editable.",
    ], [
        "Container OD: Design document “Container OD” → Table 1 “Diameter of Battery”.",
        "Container ID: Design document “Container ID”, else OD − 2 × wall (Table 2).",
        "Cathode dia: Design document “Cathode Dia” → PID electrode table (cathode row).",
        "PCD: Table 1 “Pin PCD”.   Number of holes: Table 1 “Number of Holes”.",
        "Pin dia: Table 1 “Diameter of the Pin”.",
        "Thickness and groove: design lookup tables (Table 4b / Table 3).",
    ]),
    ("TIE WIRE", [
        "Length = container height + 30.",
        "Width — by container OD: 30–70 → 3.0; above 70 → 6.0.",
        "Thickness = 0.30 (standard).",
    ], [
        "Container OD: Design document “Container OD” → Table 1 “Diameter of Battery”.",
        "Container height: Table 1 “Height of Container”.",
        "Width / thickness: tie-wire lookup table; all three overridable in CAD Revision.",
    ]),
    ("TEFLON DISC", [
        "Disc dia = lid dia − 2, where lid dia = Container ID − 0.05.",
        "Thickness = 0.20 (standard).",
        "Inner circle dia = stack dia − inner space; inner space by cathode dia: 20–30 → 3; 31–60 → 5; 61–90 → 7.",
        "Number of slots = number of tie wires; slot spacing = 360° / number of tie wires.",
        "Slots start at stack dia + 5 and run outward; slot length = (disc dia − that start dia) ÷ 2 (derived).",
        "Slot width = tie-wire width + 1.",
        "Pin holes: N holes on the PCD, spaced 360° / N; hole dia = pin dia.",
    ], [
        "Lid dia: derived from the container ID (Design document / Table 2).",
        "Cathode (stack) dia: Design document “Cathode Dia” → PID electrode table.",
        "Number of tie wires: PID configuration data — “tie wire”.",
        "PCD: Table 1 “Pin PCD”.   Number of holes: Table 1 “Number of Holes”.",
        "Hole (pin) dia: Table 1 “Diameter of the Pin”.",
    ]),
    # ---- Pellets & discs --------------------------------------------------
    ("CATHODE PELLET", [
        "Diameter = cathode diameter.",
        "Thickness = cathode electrode thickness.",
        PELLET_TOL,
    ], [
        "Dia: PID electrode table (cathode row) → Design document “Cathode Dia”.",
        "Thickness: PID electrode table, cathode row thickness.",
    ]),
    ("ANODE PELLET", [
        "Diameter = anode diameter.",
        "Thickness = anode electrode thickness.",
        PELLET_TOL,
    ], [
        "Dia: PID electrode table (anode row) → Design document “Anode Dia”.",
        "Thickness: PID electrode table, anode row thickness.",
    ]),
    ("ELECTROLYTE PELLET", [
        "Diameter = cathode diameter.",
        "Thickness = electrolyte thickness.",
        PELLET_TOL,
    ], [
        "Dia: PID electrode table (cathode row) → Design document “Cathode Dia”.",
        "Thickness: PID electrode table, electrolyte row.",
    ]),
    ("HEAT PELLET - 1", [
        "Diameter = cathode diameter.",
        "Thickness = heat pellet-1 thickness.",
        PELLET_TOL,
    ], [
        "Dia: PID electrode table (cathode row) → Design document “Cathode Dia”.",
        "Thickness: PID electrode table, “Heat Pellet-1” row.",
    ]),
    ("HEAT PELLET - 2", [
        "Diameter = cathode diameter.",
        "Thickness = heat pellet-2 thickness.",
        PELLET_TOL,
    ], [
        "Dia: PID electrode table (cathode row) → Design document “Cathode Dia”.",
        "Thickness: PID electrode table, “Heat Pellet-2” row.",
    ]),
    ("HEAT PELLET - 3", [
        "Diameter = cathode diameter.",
        "Thickness = heat pellet-3 thickness.",
        PELLET_TOL,
    ], [
        "Dia: PID electrode table (cathode row) → Design document “Cathode Dia”.",
        "Thickness: PID electrode table, “Heat Pellet-3” row.",
    ]),
    ("HEAT PELLET - 1B", [
        "Diameter = cathode diameter.",
        "Thickness = heat pellet-1B thickness.",
        PELLET_TOL,
    ], [
        "Dia: PID electrode table (cathode row) → Design document “Cathode Dia”.",
        "Thickness: PID electrode table, “Heat Pellet-1B” row.",
    ]),
    ("SS DISC", [
        "Diameter = cathode diameter.",
        "Thickness = the S.S. disc thickness given for this battery.",
        "Material SS 304.  " + PELLET_TOL,
    ], [
        "Dia: PID electrode table (cathode row) → Design document “Cathode Dia”.",
        "Thickness: PID configuration data — “S.S Disc top”.",
    ]),
    ("SS DISC (0.05)", [
        "Diameter = cathode diameter.",
        "Thickness = 0.05 (fixed standard).",
        "Material SS 304.  " + PELLET_TOL,
    ], [
        "Dia: PID electrode table (cathode row) → Design document “Cathode Dia”.",
        "Thickness: fixed in the generator; editable in CAD Revision.",
    ]),
    ("MICA DISC", [
        "Diameter = cathode diameter.",
        "Thickness = 0.15 (fixed standard).",
        PELLET_TOL,
    ], [
        "Dia: PID electrode table (cathode row) → Design document “Cathode Dia”.",
        "Thickness: fixed in the generator; editable in CAD Revision.",
    ]),
    ("SAMICA DISC", [
        "Diameter = Container ID.",
        "Thickness = 0.10 (fixed standard).",
        PELLET_TOL,
    ], [
        "Dia: Design document “Container ID”, else Container OD − 2 × wall (Table 2).",
        "Thickness: fixed in the generator; editable in CAD Revision.",
    ]),
    ("SILICON BONDED MICA DISC", [
        "Diameter = Container ID.",
        "Thickness = 1.0 (fixed standard).",
        PELLET_TOL,
    ], [
        "Dia: Design document “Container ID”, else Container OD − 2 × wall (Table 2).",
        "Thickness: fixed in the generator; editable in CAD Revision.",
    ]),
    ("MICA DISC (HOLES)", [
        "Disc dia = stack (cathode) dia.",
        "Thickness = 0.15 (standard).",
        "PCD = Pin PCD; if absent, stack dia × 0.55.",
        "Number of holes (default 6), spaced 360° / N; hole start angle 90°.",
        "Hole dia = pin dia (default 2.0).",
        PELLET_TOL,
    ], [
        "Stack dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "PCD: Table 1 “Pin PCD”.   Number of holes: Table 1 “Number of Holes”.",
        "Hole dia: Table 1 “Diameter of the Pin”.",
    ]),
    # ---- Housings ---------------------------------------------------------
    ("HOUSING - A", [
        "Outer dia = stack (cathode) dia.",
        "Inner dia = Pin PCD + 2.",
        "Thickness = mica thickness + 2 × silicon thickness (0.15 + 2 × 1.0 by default).",
        "Two cuts, each of width = squib width + 2 (squib width 5.0 default).",
        "Dia tolerances: outer +0.0 / −0.2, inner +0.2 / −0.0.",
        "BOM: Mica Ring × 1, Silicon Bonded Mica Ring (Housing-A) × 2, FiberFrax Disc × 2.",
    ], [
        "Outer dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Inner dia: Table 1 “Pin PCD” + 2 (else stack dia × 0.55 + 2).",
        "Mica / silicon thickness and squib width: standard values, editable in CAD Revision.",
        "BOM reference drawing numbers: read from those components once they are generated.",
    ]),
    ("HOUSING - B", [
        "Outer dia = Container ID − 1.",
        "Inner dia = stack (cathode) dia − 2.",
        "Thickness = 1.0 (standard).",
        "BOM: Silicon Bonded Mica Ring (Housing-B) × 1, FiberFrax Disc × 1.",
    ], [
        "Container ID: Design document “Container ID”, else Container OD − 2 × wall (Table 2).",
        "Stack dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "BOM reference drawing numbers: read from those components once they are generated.",
    ]),
    ("SILICON BONDED MICA RING (HOUSING - A)", [
        "Outer dia = stack (cathode) dia.",
        "Inner dia = Pin PCD + 2.",
        "Thickness = 1.0 (standard).",
        "Two cuts, width = squib width + 2 (squib width 5.0 default).",
        "Dia tolerances: outer +0.0 / −0.2, inner +0.2 / −0.0.",
    ], [
        "Outer dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Inner dia: Table 1 “Pin PCD” + 2 (else stack dia × 0.55 + 2).",
    ]),
    ("SILICON BONDED MICA RING (HOUSING - B)", [
        "Outer dia = Container ID − 1.",
        "Inner dia = stack (cathode) dia − 2.",
        "Thickness = 1.0 (standard).",
        "Dia tolerances: outer +0.0 / −0.2, inner +0.2 / −0.0.",
    ], [
        "Container ID: Design document “Container ID”, else Container OD − 2 × wall (Table 2).",
        "Stack dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
    ]),
    ("MICA RING", [
        "Outer dia = stack (cathode) dia.",
        "Inner dia = Pin PCD + 2.",
        "Thickness = 0.15 (standard).",
    ], [
        "Outer dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Inner dia: Table 1 “Pin PCD” + 2 (else stack dia × 0.55 + 2).",
    ]),
    # ---- Wicks & strips ---------------------------------------------------
    ("PYRO WICK - 01", [
        "Length = stack (cathode) dia + 10.",
        "Width and thickness — by container OD: up to 70 → 3.0 wide / 0.3 thick; above 70 → 6.0 / 0.7.",
    ], [
        "Stack dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Container OD: Design document “Container OD” → Table 1 “Diameter of Battery”.",
    ]),
    ("PYRO WICK - 02", [
        "Length = container height + 20.",
        "Width and thickness — by container OD: up to 70 → 3.0 / 0.3; above 70 → 6.0 / 0.7.",
    ], [
        "Container height: Table 1 “Height of Container” → Design document “Container Height”.",
        "Container OD: Design document “Container OD” → Table 1 “Diameter of Battery”.",
    ]),
    ("SAMICA STRIP", [
        "Length = container height + 20.",
        "Width = 6.0 (STD).  Thickness = 0.1 (STD).",
        "Quantity 02.",
    ], [
        "Container height: Table 1 “Height of Container” → Design document “Container Height”.",
        "Width / thickness: standard values, editable in CAD Revision.",
    ]),
    ("MICA STRIP", [
        "Length = battery height + 20.",
        "Width = 10.0 (STD).  Thickness = 0.15 (STD).",
        "Quantity 06.  Visual criteria printed on the sheet.",
    ], [
        "Battery height: PID technical specification “Dimensions (Dia × Ht)” — the height figure; "
        "falls back to Table 1 “Height of Container”.",
        "Width / thickness: standard values, editable in CAD Revision.",
    ]),
    ("GLASS CLOTH TAPE", [
        "Length = (container height ÷ tape width) × π × (stack dia + 2 × FiberFrax wrap thickness).",
        "Width = 25 (standard).  Thickness = 0.2 (standard).",
        "FiberFrax wrap thickness taken as 1.0 unless set.",
    ], [
        "Stack dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Container height: Table 1 “Height of Container” → Design document “Container Height”.",
        "Width / thickness / wrap thickness: standard values, editable in CAD Revision.",
    ]),
    ("ADHESIVE TAPE", [
        "Length 100, width 12.5, thickness 0.2 — a stock consumable, so all three are "
        "standard values, not derived from the battery.",
    ], [
        "Not battery-derived. Change any of the three in CAD Revision if a different roll is used.",
    ]),
    ("SAMICA WRAP", [
        "Length = π × (stack dia + 2 × FiberFrax wrap thickness) + 10.",
        "Width = container height − 3.",
        "Thickness = 0.1 (standard).",
    ], [
        "Stack dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Container height: Table 1 “Height of Container” → Design document “Container Height”.",
        "FiberFrax wrap thickness: 1.0 unless set in CAD Revision.",
    ]),
    ("MICA WRAP", [
        "Length = π × Container ID + 10.",
        "Width = container height − 3.",
        "Thickness = 0.1 (standard).",
    ], [
        "Container ID: Design document “Container ID”, else Container OD − 2 × wall (Table 2).",
        "Container height: Table 1 “Height of Container” → Design document “Container Height”.",
    ]),
    # ---- FiberFrax --------------------------------------------------------
    ("FIBERFRAX STACK WRAP", [
        "Length = stack circumference = π × stack (cathode) dia.",
        "Width = container height − 3.",
        "Thickness — by container OD: up to 70 → 1.0; above 70 → 1.6.",
    ], [
        "Stack dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Container height: Table 1 “Height of Container” → Design document “Container Height”.",
        "Container OD (for the thickness band): Design document “Container OD” → Table 1 “Diameter of Battery”.",
    ]),
    ("FIBERFRAX CONTAINER INSULATION", [
        "Base length = inner circumference = π × Container ID.",
        "Width = container height − 3.",
        "Thickness — by container OD: up to 70 → 1.0; above 70 → 1.6.",
        "Quantity 02 by default.",
    ], [
        "Container ID: Design document “Container ID”, else Container OD − 2 × wall (Table 2).",
        "Container height: Table 1 “Height of Container” → Design document “Container Height”.",
        "Container OD (for the thickness band): Design document “Container OD” → Table 1 “Diameter of Battery”.",
    ]),
    # ---- Terminals & misc -------------------------------------------------
    ("SQUIB TERMINAL", [
        "Length 50, width 3.0, thickness 0.2 — a bought-in nickel strip, so these are standard values.",
        "Width tolerance ±0.4; thickness tolerance +0.0 / −0.02.",
        "Material NICKEL.",
    ], [
        "Not battery-derived. Editable in CAD Revision if the supplied strip differs.",
    ]),
    ("SQUIB", [
        "A bought-in standard part — every dimension is a standard value, all editable in CAD Revision.",
        "Type is selected by the user: Single Head, or Single Head with Wired (both titled SQUIB).",
        "Single Head: overall length 12.5; head 4.0 long × 3.8 wide × 2.23 thick; strip 3.9 wide × 0.7 thick; "
        "strip length = overall − head (derived); resistance 0.80–1.20 Ω.",
        "Single Head with Wired: body height 11 (dome to base), charge height 8, body width 4 ±0.3, "
        "base width 3.8 ±0.2, body depth 3.2 ±0.2; leads 80 ±5 long × Ø0.6, spacing 1.5, "
        "lead span 1.4 inner / 1.8 outer; charge pellet 3.4 top / 3.6 bottom.",
    ], [
        "Not battery-derived — the squib is procured to a fixed specification.",
        "The type selected here drives the cut sizes on both 2-Cut discs, so generate the Squib first.",
    ]),
    ("SILICON BONDED MICA DISC (2 CUTS)", [
        "Disc dia = cathode (stack) dia.",
        "Two cuts, 360° / 2 = 180° apart; cut start angle 90°.",
        "Cut length = squib length + 2 mm clearance.",
        "Cut width = squib width + 2 mm clearance.",
        "The gap across the disc between the two cuts = disc dia − 2 × cut length (derived and dimensioned).",
        "Thickness = 1.0 (standard).  Dia tolerance +0.00 / −0.20.",
    ], [
        "Disc dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Squib size: read from the generated SQUIB drawing every time — Single Head uses its total "
        "length and head width; Single Head with Wired uses the body height (leads excluded) and body width.",
        "Clearance: 2 mm default, editable in CAD Revision.",
    ]),
    ("FIBERFRAX DISC (2 CUTS)", [
        "Identical rules to the Silicon Bonded Mica Disc (2 Cuts) above.",
        "Disc dia = cathode (stack) dia; two cuts 180° apart.",
        "Cut length = squib length + 2; cut width = squib width + 2.",
        "Thickness = 1.6 (standard).  Material FIBERFRAX.",
    ], [
        "Disc dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Squib size: read from the generated SQUIB drawing (same rule as the mica disc).",
    ]),
    ("LID WITH TIE WIRE", [
        "Lid OD, lid thickness, PCD, number of holes and hole dia — taken from the generated LID BLANK drawing.",
        "Tie-wire width and thickness — taken from the generated TIE WIRE drawing.",
        "Number of tie wires (default 3); wires equally spaced at 360° / n.",
        "Wire start dia = stack dia − 2 × 5, i.e. the wires start 5 mm radially inside the stack diameter.",
        "The wire pattern is rotated automatically to sit as far off the terminal holes as possible; "
        "a warning is raised if a wire still passes too close to a hole.",
        "Groove circle dia = PCD + 10.  Spot-weld length 5, weld strength 25 kgf.",
        "BOM: LID (or LID BLANK) and TIE WIRE, quoting their generated drawing numbers.",
    ], [
        "Lid geometry: the LID BLANK component generated for this battery (falls back to the same "
        "rules that drawing uses if it has not been generated yet).",
        "Wire width / thickness: the TIE WIRE component; else width by container OD table, thickness 0.3.",
        "Stack dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Number of tie wires: PID configuration data — “tie wire”.",
    ]),
    ("CURRENT COLLECTOR - ANODE", [
        "Disc dia = cathode diameter.",
        "Disc thickness 0.1, disc tolerance 0.05.",
        "Nickel lead: 115 long × 6 wide × 0.1 thick; gap 2.0 between disc and lead.",
        "Lead labelled NICKEL LEAD-A.  Material SS 304 / Nickel.",
    ], [
        "Disc dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Lead dimensions: standard values matching the reference Lead drawing; editable in CAD Revision.",
    ]),
    ("CURRENT COLLECTOR - CATHODE", [
        "Disc dia = cathode diameter.",
        "Disc thickness 0.1, disc tolerance 0.05.",
        "Nickel lead: 115 long × 6 wide × 0.1 thick; gap 2.0 between disc and lead.",
        "Lead labelled NICKEL LEAD-C.  Material SS 304 / Nickel.",
    ], [
        "Disc dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Lead dimensions: standard values matching the reference Lead drawing; editable in CAD Revision.",
    ]),
    ("BRACE PLATE", [
        "Outer circle dia = cathode (stack) dia.",
        "Inner circle dia = cathode dia − 2 × radial clearance (clearance 6.0 default).",
        "Number of arms = number of tie wires; arms spaced 360° / n.",
        "Small or large battery decided by container OD: 70 and below → small.",
        "Plate thickness = 0.5 (small) or 1.0 (large).",
        "Plate width = tie-wire width × 3 (small) or × 2 (large).",
        "Bump width = tie-wire width + 2; bump height = tie-wire thickness × 4; bump radius 3.0.",
        "Total section height = 2 × plate thickness + bump height; strip width = that total height; "
        "bump width in plan = tie-wire width ÷ 2.",
        "Material MS, zinc plated 12.5 microns min.; quantity = number of tie wires.",
    ], [
        "Cathode dia: PID electrode table (cathode) → Design document “Cathode Dia”.",
        "Container OD (small / large): Design document “Container OD” → Table 1 “Diameter of Battery”.",
        "Tie-wire width: tie-wire table by container OD; thickness 0.3.",
        "Number of tie wires: PID configuration “tie wire” → Table 1 “Number of Holes” → 3.",
    ]),
    ("DELIVER PIN", [
        "Pin diameter.",
        "Total pin length = Upper part of the pin + Lid blank thickness + Bottom side of the lid "
        "(recomputed on every render, never frozen).",
        "Lid blank thickness — Table 4(b) by container OD.",
        "Dia tolerance ±0.1.  Material SS 304.",
    ], [
        "Pin dia: Table 1 “Diameter of the Pin” → Design document “Pin Dia”.",
        "Upper part: Table 1 “Upper Part of the Pin”.",
        "Lid thickness: Table 4(b) using the container OD.",
        "Bottom side: entered in CAD Revision (0 if not set).",
    ]),
    # ---- Assemblies -------------------------------------------------------
    ("CELL ASSEMBLY", [
        "Six fixed rows; for each the user enters Qty/Nos and the Placing Order "
        "(the stack positions that piece occupies, e.g. “2,4”).",
        "Specified bottom (S.No 1) to top; the sheet is drawn top to bottom.",
        "Each piece is drawn at the diameter and thickness of its own generated drawing.",
        "BOM reference drawing numbers matched to the generated components by name.",
    ], [
        "Quantities and placing order: entered by the user in the CAD Drawing module and saved on the battery.",
        "Piece dia / thickness: the geometry stored on that component’s generated drawing; "
        "falls back to the cathode dia and 1.0 mm if it has not been drawn yet.",
    ]),
    ("TOP ASSEMBLY", [
        "Rows, quantities and placing order entered by the user.",
        "Drawn in the order given, with S.No balloons shown.",
        "Each piece takes the diameter and thickness of its own generated drawing.",
        "BOM reference drawing numbers matched to the generated components by name.",
    ], [
        "Quantities and placing order: entered by the user in the CAD Drawing module and saved on the battery.",
        "Piece dia / thickness: geometry of the generated component drawings.",
    ]),
    ("BOTTOM ASSEMBLY", [
        "Rows, quantities and placing order entered by the user.",
        "Drawn in the order given, with S.No balloons shown.",
        "Each piece takes the diameter and thickness of its own generated drawing.",
        "BOM reference drawing numbers matched to the generated components by name.",
    ], [
        "Quantities and placing order: entered by the user in the CAD Drawing module and saved on the battery.",
        "Piece dia / thickness: geometry of the generated component drawings.",
    ]),
    ("STACK", [
        "N = total number of cells in the stack.",
        "S = number of stacks in parallel (default 1).",
        "Quantity on the sheet = number of stacks.",
        "The cells are drawn as banded blocks with an ellipsis where the run is too long to show every cell.",
    ], [
        "Number of cells: PID configuration data — “number of cells” / “no of cells” / “total cells” / "
        "“cells” → Table 1 “Number of Cells”. Refuses to draw (400) if not found.",
        "Number of stacks: PID configuration data — “stack” or “parallel”.",
    ]),
    ("LID  (lid blank + deliver pins + G.M. seal)", [
        "Lid OD, lid thickness, PCD, number of holes, hole dia, groove depth / width, weld space and "
        "edge angle — all taken from the generated LID BLANK drawing.",
        "Pin dia, upper part and bottom side — taken from the generated DELIVER PIN drawing.",
        "Total pin length = upper part + lid thickness + bottom side, recomputed on every render so a "
        "revision to either child drawing flows straight through.",
        "G.M. seal annulus width = (hole dia − pin dia) ÷ 2; a warning is raised if the pin is not "
        "smaller than the hole.",
        "Note on the sheet: “To be filled with G.M seal at N places”, N = number of holes.",
        "Tolerances: lid OD +0.05 / −0.15, pin dia ±0.1, total pin length ±0.2.",
        "BOM: LID BLANK and DELIVER PIN with their generated drawing numbers.",
    ], [
        "Every dimension: the LID BLANK and DELIVER PIN components generated for this battery.",
        "Anything those drawings do not supply falls back to the same battery-data rules they use "
        "(Design document → PID → Table 1 → lookup tables).",
        "Generate LID BLANK and DELIVER PIN before this drawing so the BOM quotes the right numbers.",
    ]),
]

# --------------------------------------------------------------------------- #
# Appendix lookup tables
# --------------------------------------------------------------------------- #
APPENDIX = [
    ("Table 2 — Container wall thickness (by container OD)",
     ["OD band (mm)", "Deep drawn", "Flanged", "Tolerance", "Min / Max"],
     [["30 – 50", "0.70", "0.70", "+0.15 / −0.05", "0.65 / 0.85"],
      ["50 – 100", "1.10", "1.00", "±0.10", "0.90 / 1.10"],
      ["100 – 150", "1.30", "1.30", "+0.15 / −0.10", "1.20 / 1.45"]]),
    ("Table 3 — Lid groove depth and width (by container OD)",
     ["OD band (mm)", "Depth", "Width"],
     [["30 – 45", "0.80", "0.70"],
      ["46 – 100", "1.00", "1.00"],
      ["101 – 150", "1.30", "1.30"]]),
    ("Table 4 — Lid diameters and blank thickness",
     ["Item", "Rule"],
     [["Lid BLANK OD", "Container ID − 0.05  (±0.05)"],
      ["Finished lid OD", "Container ID − 0.15  (±0.05)"],
      ["Blank thickness, OD 30 – 50", "2.5"],
      ["Blank thickness, OD 51 – 80", "4.0"],
      ["Blank thickness, OD 81 – 110", "5.0"],
      ["Blank thickness, OD 111 – 150", "6.5"]]),
    ("Consumable tables (by container OD, unless stated)",
     ["Item", "Rule"],
     [["Tie-wire width", "OD 30 – 70 → 3.0;  above 70 → 6.0"],
      ["Pyro wick width / thickness", "OD up to 70 → 3.0 / 0.3;  above 70 → 6.0 / 0.7"],
      ["FiberFrax sheet thickness", "OD up to 70 → 1.0;  above 70 → 1.6"],
      ["Teflon disc inner space", "By cathode dia: 20 – 30 → 3;  31 – 60 → 5;  61 – 90 → 7"],
      ["Tie-wire length / thickness", "Container height + 30;  thickness 0.30"]]),
]


# --------------------------------------------------------------------------- #
# Document helpers
# --------------------------------------------------------------------------- #
def _shade(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _repeat_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    trPr.append(trPr.makeelement(qn("w:tblHeader"), {}))


def _lines(cell, lines: list[str], size: float = 8.5, bullet: bool = True) -> None:
    """Write one paragraph per line into a cell, replacing its empty first one."""
    cell.text = ""
    for i, txt in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(2 if i < len(lines) - 1 else 0)
        run = para.add_run(("•  " if bullet else "") + txt)
        run.font.size = Pt(size)


def build() -> Path:
    doc = Document()

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11.69), Inches(8.27)
    sec.left_margin = sec.right_margin = Inches(0.5)
    sec.top_margin = sec.bottom_margin = Inches(0.55)

    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(9)

    # ---- title page block -------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("RENEWABLE ENERGY SYSTEMS LIMITED")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RES_BLUE

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("CAD Drawing Dimension Guidelines")
    rs.bold = True
    rs.font.size = Pt(19)

    s2 = doc.add_paragraph()
    s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s2.add_run("Dimensions used by each generated component drawing, and where every value is taken from")
    r2.italic = True
    r2.font.size = Pt(10)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(10)
    ri = intro.add_run(
        "The CAD-BOM application generates 46 component drawings from the data held against a battery. "
        "This document lists, for each drawing, the dimensions it uses and the source each value is "
        "resolved from. Unless a row says otherwise, every value can be overridden per battery in the "
        "CAD Revision module."
    )
    ri.font.size = Pt(9.5)

    ph = doc.add_paragraph()
    ph.paragraph_format.space_before = Pt(8)
    rp = ph.add_run("Order in which a value is resolved")
    rp.bold = True
    rp.font.size = Pt(11)

    order = doc.add_paragraph()
    ro = order.add_run(
        "1.  Value entered by the user (CAD Revision / the drawing form)  →  "
        "2.  PID — electrode table, configuration data, technical specification  →  "
        "3.  Design document parameters  →  "
        "4.  Table 1 of the battery data  →  "
        "5.  Design lookup tables (Tables 2, 3, 4 and the consumable tables in the appendix)  →  "
        "6.  Standard value built into the generator."
    )
    ro.font.size = Pt(9.5)

    note = doc.add_paragraph()
    rn = note.add_run(
        "If a required dimension cannot be resolved at any of those steps the drawing is not produced — "
        "the application returns a message naming the missing value so it can be entered. Values shown as "
        "“derived” (container ID, total pin length, slot lengths, cut sizes, the G.M. seal annulus) are "
        "never stored: they are recalculated on every render, so revising a parent drawing flows through "
        "to everything built on it."
    )
    rn.font.size = Pt(9.5)
    rn.italic = True

    doc.add_page_break()

    # ---- main table -------------------------------------------------------
    h = doc.add_paragraph()
    rh = h.add_run("Component dimension guidelines")
    rh.bold = True
    rh.font.size = Pt(13)
    rh.font.color.rgb = RES_BLUE

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    widths = [Inches(0.45), Inches(1.85), Inches(4.35), Inches(4.05)]
    headers = ["S.No", "Component Name", "Guidelines", "Where to take that value"]

    hdr = tbl.rows[0]
    _repeat_header(hdr)
    for i, (cell, text) in enumerate(zip(hdr.cells, headers)):
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade(cell, HDR_FILL)
        cell.width = widths[i]

    for idx, (name, guide, source) in enumerate(COMPONENTS, start=1):
        row = tbl.add_row()
        c0, c1, c2, c3 = row.cells
        for cell, wdt in zip(row.cells, widths):
            cell.width = wdt

        c0.text = ""
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(idx))
        r0.bold = True
        r0.font.size = Pt(9)

        c1.text = ""
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(name)
        r1.bold = True
        r1.font.size = Pt(9)

        _lines(c2, guide)
        _lines(c3, source)

        if idx % 2 == 0:
            for cell in row.cells:
                _shade(cell, BAND_FILL)

    doc.add_page_break()

    # ---- appendix ---------------------------------------------------------
    ah = doc.add_paragraph()
    rah = ah.add_run("Appendix — design lookup tables")
    rah.bold = True
    rah.font.size = Pt(13)
    rah.font.color.rgb = RES_BLUE

    ap = doc.add_paragraph()
    rap = ap.add_run(
        "These are the tables the generator consults when a dimension is not given for the battery. "
        "They live in backend/cad/tables.py."
    )
    rap.font.size = Pt(9.5)
    rap.italic = True

    for title, cols, rows in APPENDIX:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(10)
        rc = cap.add_run(title)
        rc.bold = True
        rc.font.size = Pt(10.5)

        at = doc.add_table(rows=1, cols=len(cols))
        at.style = "Table Grid"
        at.alignment = WD_TABLE_ALIGNMENT.LEFT
        for cell, text in zip(at.rows[0].cells, cols):
            cell.text = ""
            run = cell.paragraphs[0].add_run(text)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _shade(cell, HDR_FILL)
        for vals in rows:
            cells = at.add_row().cells
            for cell, text in zip(cells, vals):
                cell.text = ""
                run = cell.paragraphs[0].add_run(text)
                run.font.size = Pt(9)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"written: {path}  ({path.stat().st_size:,} bytes)")
